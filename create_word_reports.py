import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    doc.add_heading('Звіт', 0)

    lines = content.split('\n')
    current_paragraph = []

    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line == '':
            if current_paragraph:
                doc.add_paragraph(' '.join(current_paragraph))
                current_paragraph = []
        else:
            current_paragraph.append(line)

    if current_paragraph:
        doc.add_paragraph(' '.join(current_paragraph))

    doc.save(docx_file)
    print(f"Created {docx_file}")

# List of reports
reports = [
    ('report_neural_network.md', 'Звіт_нейронні_мережі.docx'),
    ('report_aiogram_bot.md', 'Звіт_Aiogram_бот.docx'),
    ('report_telebot.md', 'Звіт_Telebot.docx'),
    ('report_web.md', 'Звіт_веб_ресурс.docx')
]

for md, docx in reports:
    md_to_docx(md, docx)